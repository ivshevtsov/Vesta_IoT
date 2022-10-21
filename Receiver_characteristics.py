import numpy as np
from docx import Document
##---------------------------------##
##          Initial data           ##
##---------------------------------##
Loss = 3

Sens = -108.2
Max_Input_Sig = -25 - Loss
ADC_FS_Vpp = 1
ADC_FS_Vrms = ADC_FS_Vpp/(2*2**0.5)
ADC_FS_dBm = 20*np.log10(ADC_FS_Vrms*1e6)-10*np.log10(50)-90

BW = 200e3
SNR = 1.5

#--Tests number
IBB_Test=1
ACS_Test=1

#out of band blockers

out_band_value = [Sens+6, -44, -30, -15]
out_band_freq_offset = [0, 15e6, 60e6, 85e6]

#in band blockers
if IBB_Test==1:
    in_band_value = [Sens + 6, -56 ]
    in_band_freq_offset = [0, 7.5e6]
    in_band_band = [200e3, 5e6]
elif IBB_Test==2:
    in_band_value = [Sens + 6, -44 ]
    in_band_freq_offset = [0, 12.5e6]
    in_band_band = [200e3, 5e6]
else:
    print('Enter correct test number for IBB(IBB_Test[1,2])')

#adjasted blockers
if ACS_Test==1:
    adj_value = [Sens+14, Sens+42, Sens+47]
    adj_freq_offset = [0, 200e3, 2.5e6]
    adj_band = [200e3, 200e3, 5e6]
elif ACS_Test==2:
    adj_value = [-53, -25, -25]
    adj_freq_offset = [0, 200e3, 2.5e6]
    adj_band = [200e3, 200e3, 5e6]
else:
    print('Enter correct test number for ACS(ACS_Test[1,2])')

#intermodulation characteristics
Pb_value = -46

doc = Document('DATA/Receiver_characteristics/Characteristics.docx')
doc.tables[0].cell(1,1).text = str(Loss)
doc.tables[0].cell(2,1).text = str(Sens)
doc.tables[0].cell(3,1).text = str(Max_Input_Sig) + '*'
doc.tables[0].cell(4,1).text = str(round(ADC_FS_dBm,2))
doc.tables[0].cell(5,1).text = str(BW/1e3)
doc.tables[0].cell(6,1).text = str(SNR)
doc.save('DATA/Receiver_characteristics/Characteristics.docx')


##---------------------------------##
##    Calculate characteristics    ##
##---------------------------------##

#receiver noise figure
receiver_noise = 174+Sens-SNR-10*np.log10(BW)-Loss

#receiver minimum/maximum gain
receiver_gain_max = ADC_FS_dBm-Sens
receiver_gain_min = ADC_FS_dBm-Max_Input_Sig
receiver_gain_dynamic = receiver_gain_max-receiver_gain_min

doc = Document('DATA/Receiver_characteristics/Characteristics.docx')
doc.tables[1].cell(1,1).text = str(round(receiver_noise,2))
doc.tables[1].cell(2,1).text = str(round(receiver_gain_max,2))
doc.tables[1].cell(3,1).text = str(round(receiver_gain_min,2))
doc.tables[1].cell(4,1).text = str(round(receiver_gain_dynamic,2))
doc.save('DATA/Receiver_characteristics/Characteristics.docx')

#Out of band attenuation
#out power must be less or equal in band power
out_of_band_att=[]
for i in range(len(out_band_value)):
    out_of_band_att.append(out_band_value[i]-Max_Input_Sig+SNR)
for i in range(len(out_band_value)):
    if out_of_band_att[i]>=0:
        print(f'1.Out attenuation {out_of_band_att[i]} dB at'
              f'freq +-{out_band_freq_offset[i]/1e6} MHz\n'
              f'from carrier frequency')

#in band blockers and characteristics
in_band_phase_noise = []
for i in range(len(in_band_value)):
    in_band_phase_noise.append(-174-in_band_value[i])
    print(f'2.In Band Phase Noise {in_band_phase_noise[i]} dBc/Hz'
          f' at freq {in_band_freq_offset[i]/1e6} Mhz')


#adjasted channel calculation
IMRR = (adj_value[1]-adj_value[0])+SNR
print(f'3.IMRR Value {IMRR} db with IF<=200kHz')

#calculate IIP3
IIP3 = (3*Pb_value-Sens+SNR)/2
print(f'4.Calculated IIP3 = {round(IIP3,2)} dBm')

#Clculate compression point
P1dB = Max_Input_Sig+3
print(f'5.Calculated P1dB = {round(P1dB,2)} dBm')
print(f'5.Calculated IIP3 = P1dB+10 = {round(P1dB+10,2)} dBm')



