from openpyxl import load_workbook
from docx import Document
from docx.shared import Cm

# ------------------------------------------------ #
# ------------------------------------------------ #


def get_key(d, value):
    string = ', '
    counter = 0
    for k, v in d.items():
        if v == value:
            counter = counter+1
            if counter == 1:
                string = str(k)
            else:
                string = string + ', ' + str(k)
    return string


def exel_dictionary(HOME, File, nsheet, nrow, ncol):
    # load document
    wb = load_workbook(f'{HOME}/{File}.xlsx')
    # get sheet names
    sheetnames = wb.sheetnames
    # get sheet data
    sheet = wb[sheetnames[nsheet - 1]]
    # list elements
    List_elements = []
    dic = {}
    with open(f'{HOME}/{File}.txt', "w") as file:
        for row in range(2, nrow + 1):
            for column in range(2, ncol + 1):
                row_x = sheet.cell(row=row, column=1).value
                column_x = sheet.cell(row=1, column=column).value
                Value = sheet.cell(row=row, column=column).value
                Coordinate = f'{row_x}' + f'{column_x}'
                String = Coordinate + ' ' + f'{Value}'
                dic[Coordinate] = f'{Value}'
                # create list elements for compare
                List_elements.append(String)
                # print(String)
                file.write(f'{String} \n')
        file.close()
    return dic


def exel_dictionary_coord(HOME, File, nsheet, nrow, ncol, pitch):
    # document
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Номер вывода кристалла'
    table.cell(0, 1).text = 'Обозначение вывода кристалла'
    table.cell(0, 2).text = 'X'
    table.cell(0, 3).text = 'Y'
    sep = '	'
    # load document
    wb = load_workbook(f'{HOME}/{File}.xlsx')
    # get sheet names
    sheetnames = wb.sheetnames
    # get sheet data
    sheet = wb[sheetnames[nsheet - 1]]
    # list elements
    List_elements = []
    dic = {}
    i = 0
    for row in range(2, nrow + 1):
        for column in range(2, ncol + 1):
            i = i + 1
            row_x = sheet.cell(row=row, column=1).value
            column_x = sheet.cell(row=1, column=column).value
            Value = sheet.cell(row=row, column=column).value
            Coordinate = f'{row_x}' + f'{column_x}'

            Coordinate_X = (-pitch * (nrow - 2) / 2) + (column - 2) * pitch
            Coordinate_Y = (pitch * (nrow - 2) / 2) - (row - 2) * pitch

            String = Coordinate + sep + f'{Value}' + sep + f'{Coordinate_X},{Coordinate_Y}'
            dic[Coordinate] = f'{Coordinate_X},{Coordinate_Y}'

            # work with doc
            # table.cell(i+1, 0).text = str(Coordinate)
            # table.cell(i+1, 1).text = str(Value)
            # table.cell(i+1, 2).text = str(Coordinate_X)
            # table.cell(i+1, 3).text = str(Coordinate_Y)
            row_cells = table.add_row().cells
            row_cells[0].text = str(Coordinate)
            row_cells[1].text = str(Value)
            row_cells[2].text = str(Coordinate_X)
            row_cells[3].text = str(Coordinate_Y)
    doc.save(f'{HOME}/{File}.docx')
    return dic


def doc_matching(HOME, File):
    doc = Document(f'{HOME}/{File}.docx')
    table_size = len(doc.tables[0].rows)

    dic_match = {}
    for i in range(table_size - 1):
        package_name = doc.tables[0].rows[i + 1].cells[0].text
        chip_name = doc.tables[0].rows[i + 1].cells[1].text

        dic_match[chip_name] = package_name
    return dic_match


# ------------------------------------------------ #
# ------------------------------------------------ #

# directory
HOME = 'DATA/EXEL_package_chip'

# create dictionaries
dic_package = exel_dictionary(HOME, File='PACKAGE_PADS', nsheet=5, nrow=22, ncol=22)
dic_chip = exel_dictionary(HOME, File='CHIP_PADS', nsheet=1, nrow=62, ncol=62)
dic_match = doc_matching(HOME, File='MATCH_LIST')


# Create  common table
document = Document()

# set list area
sections = document.sections
for section in sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)


# create initial state for table
table = document.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.cell(0, 0) .text = 'Номер вывода кристалла'
table.cell(0, 1) .text = 'Обозначение вывода крсталла'
table.cell(0, 2) .text = 'Номер вывода корпуса'
table.cell(0, 3) .text = 'Обозначение вывода корпуса'


for key in dic_chip:

    # chip coordinates
    chip_coordinate = key
    chip_name = dic_chip[key]

    if chip_name != 'no bump':

        row_cells = table.add_row().cells

        row_cells[0].text = chip_coordinate
        row_cells[1].text = chip_name

        # package coordinates
        if chip_name in dic_match:
            package_name = dic_match[chip_name]
            package_coordinate = get_key(dic_package, package_name)

            row_cells[2].text = package_coordinate
            row_cells[3].text = package_name

        else:
            if chip_name != 'NC':
                print(chip_name)


document.save(f'{HOME}/TABLE.docx')


'''

#read doc description
doc=Document(f'{HOME}/DESCRIPTION.docx')
table_size = len(doc.tables[0].rows)

dic_discr={}
for i in range(table_size-1):
    name =        doc.tables[0].rows[i+1].cells[0].text
    description = doc.tables[0].rows[i+1].cells[1].text
    dic_discr[name]=description


for key in dic_package:
    row_cells = table.add_row().cells
    row_cells[0].text = key
    row_cells[1].text = dic_package[key]
    row_cells[2].text = get_key(dic_chip, dic_package[key])
    row_cells[3].text = dic_discr[dic_package[key]]

document.save(f'{HOME}/Table.docx')
'''