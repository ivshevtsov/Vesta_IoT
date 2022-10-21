from docx import Document

HOME = 'DATA/EXEL_convert'
File = 'RAYAZH'

doc = Document(f'{HOME}/{File}.docx')
table = doc.tables[0]
all_tables = doc.tables
print('Всего таблиц в документе:', len(all_tables))

print(table.cell(0,2).text)
